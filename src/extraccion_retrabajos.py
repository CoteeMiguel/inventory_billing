import docx
import os
import openpyxl
import os
from dotenv import load_dotenv
load_dotenv()

class WordTableExtractor:
    def __init__(self):
        self.outputfile = os.getenv('PATH_OUTPUT')

    def extractfromfolder(self,ruta):

        archivo = open(self.outputfile,'w')
        listaretrabajos = os.listdir(ruta)
        #print(listaretrabajos)
        ListaRegistro = []
        ListaItem = [] 
        for solicitud in listaretrabajos:
            rutasolicitud = ruta + "\\" + solicitud
            documento = docx.Document(rutasolicitud)
            tables = documento.tables
            #ListaItem.append(solicitud)
            archivo.write(solicitud)
            archivo.write(";")
            
            """Para iterar sobre  la tabla del documento ingresado"""
            for table in tables:
                for row in table.rows:
                    for cell in row.cells:  
                        for paragraph in cell.paragraphs:
                            #print(paragraph.text)
                            #ListaItem.append(paragraph.text)
                            archivo.write(paragraph.text + ";")
            #ListaRegistro.append(ListaItem)
            archivo.write("\n")

if __name__ == '__main__':
    ruta = input("Cuál es la carpeta?")
    WordTableExtractor.extractfromfolder()
        